__author__ = 'abigail'
__purpose__ = 'Parse and copy records to EDC'


"""
This script parses and copies info from word docs to EDC for brain MRI 
"""

import os
from redcap import Project 
from datetime import datetime
from docx import Document
import re
import docx2txt
import subprocess
import pandas as pd

#set redcap variables
serviceurl = 'https://redcap.vanderbilt.edu/api/'
#edc_database = Project(serviceurl, os.environ['VMAP_ELECTRONIC_DATA_CAPTURE'])
#pull redcap data for all brain MRI reports
redcap_data = edc_database.export_records(fields=['map_id', 'session_id'])
staff = Project(serviceurl, os.environ['STAFF_COMPLIANCE'])
staff_data = staff.export_records(fields=['staff_name'])


def get_cmd_line():
    """
    Options to run for either 1 record, or all batches of records
    :return:
    """
    from argparse import ArgumentParser
    parser = ArgumentParser()
    parser.add_argument("-p", "--path", type=str, nargs='+', help="Ids to move information for")
    args = parser.parse_args()
    return args.path


def get_staff_id(staff_name):
    try:
        return [item for item in staff_data if item['staff_name'] == staff_name][0]['staff_id']
    except IndexError:
        return None


def get_tech_name(tech_name):
    tech_list = {'Chris Thompson': '1', 'Clair Jones':  '2', 'Cori Welliever': '3', 'Emily Patterson': '7',
                 'Josh Hageman': '4', 'Marisa Bush': '5', 'Rhiannon Jones': '6', 'Stephanie Ervin': '8', 'David Pennell': '9'}
    return tech_list[tech_name]


if __name__ == '__main__':
    # Mapping the visit names to epochs
    visit_epoch = {'Enrollment': '1', '18 Month Follow-Up': '2', '36 Month Follow-Up': '3', '60 Month Follow-Up': '4',
               '7 Year Follow-Up': '5'}
    # Mapping scan names to  0: the redcap scan_acquired variable prefixes, 1: the specific scan description on XNAT
    scan_names = {'Scout': 'scout', 'VEASL': 'asl', 'SWI': 'swi', 'M0': 'mo', '3DT1W': 't1', 'T2 FLAIR': 'flair', 'DTI': 'dti',
              'B0 Field Map': 'bo', 'TRUST': 'trust', 'PCA': 'pca', 'Vessel Wall': 'vwi', 'FE_EPI_32chSHC': 'fmri',
              'PC_': 'neck', 'MJD Neck_ves' : 'mjd_neck_ves', 'SWI' : 'swi'}
    # Possible comments  for when a scan is not collected
    no_scan = ['not collected', 'not acquired', 'acquired last time', 'already acquired', 'not run', '']

    #Make an empty dictionary to add info to 
    to_upload = []

    #look at each path in paths
    paths = get_cmd_line()
    for p in paths:
        #change directory to path
        os.chdir(p)
        #for each document in the path
        for document_name in [item for item in os.listdir(p) if item.endswith('docx') and not item.startswith('~')]:
            report = Document('{}/{}'.format(p, document_name))
            # Holder for output that can be uploaded to REDCap
            nil_output = {}
            """ Extract PTP information from the document """
            # To get session ID, we need to extract the paragraph text from the document
            para_text = [item.text for item in report.paragraphs]
            # Get the session ID field
            session_field = [text for text in para_text if 'scan id' in text.lower()]
            if session_field:
                sess = session_field[0].split(":")[1]
                try:
                    session_id = sess.strip().split()[0]
                except IndexError:
                    print('{} is the sess in {}'.format(sess, document_name))
                # Find matching redcap dict
                nil_match = [item for item in redcap_data if session_id == item['session_id']]
                if nil_match:
                    nil_output = nil_match[0]
                    #print(nil_output)
                    # MAP ID, scan date and visit information is in the first table of the document
                    # Extract all the rows from this table
                    rows = report.tables[0].rows
                    # Extract all cell data from these rows
                    all_cell_data = []
                    for row in rows:
                        all_cell_data.append([c.text for c in row.cells])
                    # Split into the 2 column format
                    cell_data = []
                    for info in all_cell_data:
                        cell_data.append(info[0:2])
                        cell_data.append(info[2:4])
                    # Now find the row that each of the elements we want belong to
                    # MAP ID is in the first row
                    map_info = [item for item in cell_data if 'Participant' in item][0]
                    assert nil_output['map_id'] == map_info[1].split()[1], "{} {}".format(document_name, session_id)
                    # Scan date is in the second row
                    scan_date_orig = [item for item in cell_data if 'Date' in item][0][1]
                    # Need to convert scan date to the format needed by REDCap
                    if scan_date_orig:
                        try:
                            scan_date_object = datetime.strptime(scan_date_orig, '%m/%d/%Y')
                        except ValueError:
                            try:
                                scan_date_object = datetime.strptime(scan_date_orig, '%d %B %Y')
                            except ValueError:
                                try:
                                    scan_date_object = datetime.strptime(scan_date_orig, '%d %b %Y')
                                except ValueError:
                                    scan_date_object = datetime.strptime(scan_date_orig, '%m/%d/%y')
                        scan_date_final = scan_date_object.strftime('%Y-%m-%d')
                    else:
                        print(document_name)
                        continue
                    #Get MAP imaging tech
                    try:
                        maptech = [item for item in cell_data if any('Imaging Tech' in s for s in item)][0][1]
                        nil_output['map_imaging_tech'] = get_staff_id(maptech)
                    except IndexError:
                        nil_output['map_imaging_tech'] = ''
                    #Get VUIIS tech
                    try:
                        vuiistech = [item for item in cell_data if 'VUIIS Technologist' in item][0][1]
                        nil_output['vuiis_tech'] = get_tech_name(vuiistech)
                    except (IndexError, KeyError) as err:
                        nil_output['vuiis_tech'] = ''
                    #get appointment time
                    time_a = [item for item in cell_data if 'Appointment Time' in item]
                    if time_a:
                        time_a = time_a[0][1].split(" ")[0].replace("*", "").replace('.', '')
                        if time_a != '':
                            try:
                                in_time_a = datetime.strptime(time_a, "%I:%M%p")
                            except ValueError:
                                if int(time_a.split(':')[0]) < 6 or int(time_a.split(':')[0]) == 12:
                                    time_a = time_a + 'PM'
                                else:
                                    time_a = time_a + 'AM'
                                in_time_a = datetime.strptime(time_a, "%I:%M%p")
                            out_time_a = datetime.strftime(in_time_a, "%H:%M")
                            nil_output['appointment_time'] = out_time_a
                    #get scanner entry time
                    time_s = [item for item in cell_data if 'Scanner Entry Time' in item]
                    if time_s:
                        time_s = time_s[0][1].split(" ")[0].replace("*", "").replace('.', '')
                        if time_s != '':
                            try:
                                in_time_s = datetime.strptime(time_s, "%I:%M%p")
                            except ValueError:
                                if int(time_s.split(':')[0]) < 6 or int(time_s.split(':')[0]) == 12:
                                    time_s = time_s + 'PM'
                                else:
                                    time_s = time_s + 'AM'
                                in_time_s = datetime.strptime(time_s, "%I:%M%p")
                            out_time_s = datetime.strftime(in_time_s, "%H:%M")
                            nil_output['scanner_entry_time']= out_time_s
                    #get end time
                    time_e = [item for item in cell_data if 'End Time' in item]
                    if time_e:
                        time_e = time_e[0][1].split(" ")[0].replace("*", "").replace('.', '')
                        if time_e != '':
                            try:
                                in_time_e = datetime.strptime(time_e, "%I:%M%p")
                            except ValueError:
                                if int(time_e.split(':')[0]) < 6 or int(time_e.split(':')[0]) == 12:
                                    time_e = time_e + 'PM'
                                else:
                                    time_e = time_e + 'AM'
                                in_time_e = datetime.strptime(time_e, "%I:%M%p")
                            out_time_e = datetime.strftime(in_time_e, "%H:%M")
                            nil_output['end_time']= out_time_e
                    #get screening time
                    time_screen = [item for item in cell_data if 'Screening Time' in item]
                    if time_screen:
                        time_screen = time_screen[0][1].split(" ")[0].replace("*", "").replace('.', '')
                        if time_screen != '':
                            try:
                                in_time_screen = datetime.strptime(time_screen, "%I:%M%p")
                            except ValueError:
                                if int(time_screen.split(':')[0]) < 6 or int(time_screen.split(':')[0]) == 12:
                                    time_screen = time_screen + 'PM'
                                else:
                                    time_screen = time_screen + 'AM'
                                in_time_screen = datetime.strptime(time_screen, "%I:%M%p")
                            out_time_screen = datetime.strftime(in_time_screen, "%H:%M")
                            nil_output['scan_date_time'] = scan_date_final + ' ' + out_time_screen
                            print(nil_output)
                    to_upload.append(nil_output)
                    to_sync = nil_output['session_id']
            #print("Syncing record for session {to_sync}".format(to_sync=to_sync))
            else:
                print('{} has no session id'.format(document_name))


        """ Upload this information to REDCap """
        edc_database.import_records(to_upload)
        #print(len(to_upload))
